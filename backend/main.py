from __future__ import annotations

import asyncio
import csv
import io
import json
import sys
import uuid
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import Depends, FastAPI, HTTPException, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel

ROOT = Path(__file__).resolve().parents[1]
BACKEND_ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from agents.common import LLMGateway, output_dir
from agents.cortana.rag import answer as cortana_answer
from agents.heimdall.router import route_request
from agents.jarvis.jira_client import create_tickets_if_configured
from agents.jarvis.ticket_generator import Ticket, generate_tickets
from app.auth import get_current_demo_user
from app.settings import settings
from app.workflows import WORKFLOWS, workflow_summaries


class StartRunRequest(BaseModel):
    workflow_id: str


class ChatRequest(BaseModel):
    message: str


class ExportRequest(BaseModel):
    format: str = "csv"
    tickets: list[dict[str, Any]]
    project_key: str = "ASGARD"


class RunState(dict):
    pass


RUNS: dict[str, RunState] = {}
CHAT: dict[str, list[dict[str, str]]] = {}


class WebSocketManager:
    def __init__(self) -> None:
        self.connections: dict[str, set[WebSocket]] = {}

    async def connect(self, run_id: str, websocket: WebSocket) -> None:
        await websocket.accept()
        self.connections.setdefault(run_id, set()).add(websocket)

    def disconnect(self, run_id: str, websocket: WebSocket) -> None:
        self.connections.get(run_id, set()).discard(websocket)

    async def broadcast(self, run_id: str, event: dict[str, Any]) -> None:
        stale: list[WebSocket] = []
        for websocket in self.connections.get(run_id, set()):
            try:
                await websocket.send_json(event)
            except RuntimeError:
                stale.append(websocket)
        for websocket in stale:
            self.disconnect(run_id, websocket)


ws_manager = WebSocketManager()


app = FastAPI(title="Asgard Agentic Workflow Platform", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://localhost:3000",
        settings.frontend_url,
        "https://asgard.buildwithayush.co.in",
        "https://www.buildwithayush.co.in",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health() -> dict[str, str]:
    return {"status": "ok", "app": "asgard"}


@app.get("/api/workflows/")
async def list_workflows(user: dict = Depends(get_current_demo_user)) -> list[dict]:
    return workflow_summaries()


@app.post("/api/runs/")
async def start_run(body: StartRunRequest, user: dict = Depends(get_current_demo_user)) -> dict:
    if body.workflow_id not in WORKFLOWS:
        raise HTTPException(status_code=404, detail="Workflow not found")

    run_id = str(uuid.uuid4())
    workflow = WORKFLOWS[body.workflow_id]
    RUNS[run_id] = RunState(
        run_id=run_id,
        workflow_id=body.workflow_id,
        status="running",
        state={},
        triggered_by=user["email"],
        created_at=datetime.utcnow().isoformat(),
    )
    CHAT[run_id] = [
        {
            "role": "assistant",
            "content": f"{workflow.persona_name} is ready. Send a prompt or try one of the sample requests.",
            "created_at": datetime.utcnow().isoformat(),
        }
    ]
    return {"run_id": run_id, "status": "running", "workflow_id": body.workflow_id}


@app.get("/api/runs/{run_id}")
async def get_run(run_id: str, user: dict = Depends(get_current_demo_user)) -> dict:
    run = RUNS.get(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    return run


@app.get("/api/runs/{run_id}/chat")
async def get_chat(run_id: str, user: dict = Depends(get_current_demo_user)) -> list[dict]:
    if run_id not in RUNS:
        raise HTTPException(status_code=404, detail="Run not found")
    return CHAT.get(run_id, [])


@app.post("/api/runs/{run_id}/chat")
async def send_chat(run_id: str, body: ChatRequest, user: dict = Depends(get_current_demo_user)) -> dict:
    run = RUNS.get(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    message = body.message.strip()
    if not message:
        raise HTTPException(status_code=400, detail="Message is required")

    user_message = {"role": "user", "content": message, "created_at": datetime.utcnow().isoformat()}
    CHAT[run_id].append(user_message)
    await ws_manager.broadcast(run_id, {"event": "chat_message", "role": "user", "content": message})
    asyncio.create_task(_process_agent_message(run_id, message))
    return {"status": "sent"}


@app.websocket("/api/runs/{run_id}/ws")
async def run_ws(websocket: WebSocket, run_id: str) -> None:
    if run_id not in RUNS:
        await websocket.close(code=1008)
        return
    await ws_manager.connect(run_id, websocket)
    try:
        for message in CHAT.get(run_id, []):
            await websocket.send_json({"event": "chat_message", **message})
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        ws_manager.disconnect(run_id, websocket)


@app.post("/api/runs/{run_id}/upload")
async def upload_file(run_id: str, file: UploadFile, user: dict = Depends(get_current_demo_user)) -> dict:
    if run_id not in RUNS:
        raise HTTPException(status_code=404, detail="Run not found")
    suffix = Path(file.filename or "upload.txt").suffix
    safe_name = f"{uuid.uuid4()}{suffix}"
    path = output_dir() / safe_name
    content = await file.read()
    path.write_bytes(content)
    return {"file_path": str(path), "filename": file.filename}


@app.post("/api/runs/{run_id}/export-now")
async def export_now(run_id: str, body: ExportRequest, user: dict = Depends(get_current_demo_user)) -> Response:
    if run_id not in RUNS:
        raise HTTPException(status_code=404, detail="Run not found")
    if body.format not in {"csv", "docx"}:
        raise HTTPException(status_code=400, detail="format must be csv or docx")
    if body.format == "docx":
        return _docx_response(body.tickets, body.project_key)
    return _csv_response(body.tickets, body.project_key)


@app.get("/api/integrations/status")
async def integration_status(user: dict = Depends(get_current_demo_user)) -> dict:
    return {
        "integrations": [
            {
                "id": "llm",
                "name": f"{settings.llm_provider.upper()} LLM",
                "status": "connected" if _provider_configured() else "not_configured",
                "last_checked": datetime.utcnow().isoformat(),
            },
            {
                "id": "atlassian",
                "name": "Atlassian",
                "status": "connected" if settings.jira_api_token else "not_configured",
                "last_checked": datetime.utcnow().isoformat(),
            },
            {
                "id": "figma",
                "name": "Figma",
                "status": "connected" if settings.figma_service_account_token else "not_configured",
                "last_checked": datetime.utcnow().isoformat(),
            },
            {
                "id": "knowledge_base",
                "name": "Sample Knowledge Base",
                "status": "connected",
                "last_checked": datetime.utcnow().isoformat(),
            },
        ]
    }


@app.post("/api/integrations/{integration_id}/reconnect")
async def reconnect_integration(integration_id: str, user: dict = Depends(get_current_demo_user)) -> dict:
    return {
        "success": False,
        "id": integration_id,
        "reconnect_unsupported": True,
        "message": "Hosted Asgard does not collect credentials. Clone the repo and configure .env to enable this integration.",
    }


async def _process_agent_message(run_id: str, message: str) -> None:
    run = RUNS[run_id]
    workflow_id = run["workflow_id"]
    await ws_manager.broadcast(
        run_id,
        {
            "event": "agent_log",
            "agent": workflow_id,
            "level": "info",
            "log_event": "message_received",
            "metadata": {"provider": settings.llm_provider},
        },
    )

    if workflow_id == "heimdall":
        await _run_heimdall(run_id, message)
    elif workflow_id == "cortana":
        await _run_cortana(run_id, message)
    elif workflow_id == "jarvis":
        await _run_jarvis(run_id, message)


async def _run_heimdall(run_id: str, message: str) -> None:
    decision = route_request(message)
    fallback = (
        f"Route decision: **{decision.agent.title()}**\n\n"
        f"Confidence: **{decision.confidence}**\n\n"
        f"Reason: {decision.reason}"
    )
    text = await LLMGateway().complete(
        "You are Heimdall, a concise PM workflow router. Return a routing decision, confidence, and reason.",
        message,
        fallback=fallback,
    )
    await _assistant(run_id, text, preview={"type": "route", **asdict(decision)})
    await ws_manager.broadcast(
        run_id,
        {
            "event": "agent_log",
            "agent": "Heimdall",
            "level": "info",
            "log_event": "routing_decision",
            "metadata": asdict(decision),
        },
    )


async def _run_cortana(run_id: str, message: str) -> None:
    result = cortana_answer(message, ROOT / "samples" / "docs")
    fallback = result["answer"]
    text = await LLMGateway().complete(
        "You are Cortana, a document Q&A agent. Answer only from supplied sample context and cite source names.",
        json.dumps({"question": message, "sources": result["sources"]}),
        fallback=fallback,
    )
    await _assistant(run_id, text, preview={"type": "answer", "content": text, "sources": result["sources"]})
    await ws_manager.broadcast(
        run_id,
        {
            "event": "agent_log",
            "agent": "Cortana",
            "level": "info",
            "log_event": "answer_sources",
            "metadata": {"sources": result["sources"]},
        },
    )


async def _run_jarvis(run_id: str, message: str) -> None:
    fallback_tickets = generate_tickets(message)
    ticket_prompt = (
        "You are Jarvis, a senior product manager. Improve the following Jira-ready ticket plan, "
        "but keep it concise and suitable for a recruiter demo."
    )
    summary = await LLMGateway().complete(
        ticket_prompt,
        json.dumps([ticket.to_dict() for ticket in fallback_tickets]),
        fallback="I prepared a Jira-ready ticket set. Review it in the Output panel, then export CSV or DOCX.",
    )
    tickets = [ticket.to_dict() for ticket in fallback_tickets]
    jira_result = create_tickets_if_configured(tickets)
    RUNS[run_id]["state"] = {"tickets": tickets, "jira": jira_result}
    await _assistant(run_id, summary, preview={"type": "tickets", "tickets": tickets, "jira": jira_result})
    await ws_manager.broadcast(
        run_id,
        {
            "event": "agent_log",
            "agent": "Jarvis",
            "level": "info",
            "log_event": "tickets_preview",
            "metadata": {"tickets": tickets, "jira": jira_result},
        },
    )


async def _assistant(run_id: str, text: str, preview: dict | None = None) -> None:
    message = {"role": "assistant", "content": text, "created_at": datetime.utcnow().isoformat()}
    CHAT[run_id].append(message)
    await ws_manager.broadcast(run_id, {"event": "chat_message", **message, "preview": preview})


def _csv_response(tickets: list[dict[str, Any]], project_key: str) -> Response:
    buffer = io.StringIO()
    writer = csv.DictWriter(
        buffer,
        fieldnames=["issue_type", "summary", "description", "acceptance_criteria", "priority", "labels", "story_points"],
    )
    writer.writeheader()
    for ticket in tickets:
        writer.writerow(
            {
                **ticket,
                "acceptance_criteria": "\n".join(ticket.get("acceptance_criteria", [])),
                "labels": ",".join(ticket.get("labels", [])),
            }
        )
    filename = f"asgard_tickets_{project_key}.csv"
    return Response(
        buffer.getvalue(),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _docx_response(tickets: list[dict[str, Any]], project_key: str) -> Response:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=503, detail="python-docx is not installed") from exc

    document = Document()
    document.add_heading(f"Asgard Jira Tickets: {project_key}", 0)
    for ticket in tickets:
        document.add_heading(f"{ticket.get('issue_type', 'Task')}: {ticket.get('summary', '')}", level=1)
        document.add_paragraph(ticket.get("description", ""))
        document.add_paragraph(f"Priority: {ticket.get('priority', 'Medium')}")
        document.add_paragraph("Acceptance criteria:")
        for criterion in ticket.get("acceptance_criteria", []):
            document.add_paragraph(criterion, style="List Bullet")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f"asgard_tickets_{project_key}.docx"
    return Response(
        buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _provider_configured() -> bool:
    if settings.llm_provider == "groq":
        return bool(settings.groq_api_key)
    if settings.llm_provider == "anthropic":
        return bool(settings.anthropic_api_key)
    return False
